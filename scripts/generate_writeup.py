"""Generate the portfolio analysis writeup as a .docx file."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def set_cell_text(cell, text, bold=False, size=9):
    cell.text = ""
    run = cell.paragraphs[0].add_run(text)
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    run.bold = bold


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, size=9)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            set_cell_text(table.rows[r_idx + 1].cells[c_idx], str(val), size=9)
    return table


doc = Document()

# -- Styles --
style = doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.space_before = Pt(0)

for level in range(1, 4):
    hs = doc.styles[f"Heading {level}"]
    hs.font.name = "Calibri"
    hs.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

# ═══════════════════════════════════════════════════════════════════════════
# TITLE
# ═══════════════════════════════════════════════════════════════════════════
title = doc.add_heading("Lending Club Loan Portfolio: Investment Analysis Writeup", level=1)

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 1 — DATA ISSUES
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading("1. Data Issues and How They Were Handled", level=2)

doc.add_paragraph(
    "The raw Lending Club dataset contains 2,260,701 loans across 151 columns, spanning "
    "vintages from 2007 Q2 through 2018 Q4, with a snapshot date of March 2019. "
    "Several data quality issues required systematic cleaning before the portfolio "
    "could be analyzed. In total, 5,207 rows (0.23%) were dropped and 122 columns "
    "were removed, with full traceability documented for each step."
)

doc.add_heading("Structural Issues", level=3)

doc.add_paragraph(
    "The CSV contained 33 trailing summary rows with non-numeric ID values that were "
    "not actual loan records. These were identified and removed by filtering on the ID column. "
    "Of the 151 raw columns, only 29 are referenced by the analytics pipeline, so a strict "
    "whitelist approach was applied, reducing the SQLite database from over 1 GB to approximately "
    "200 MB without losing any analytical capability."
)

doc.add_heading("Missing and Anomalous Data", level=3)

doc.add_paragraph(
    "2,427 loans had no last payment date (last_pymnt_d), making it impossible to "
    "compute amortization schedules or identify the cash flow population. These were dropped. "
    "An additional 2,737 loans (0.12%) carried a \"Does not meet the credit policy\" status, "
    "representing legacy originations under a defunct credit framework. Since these are not "
    "representative of the investable universe as of 2019, they were excluded. Six Current loans "
    "had last payment dates before February 2019, indicating stale or anomalous records, and "
    "were also removed."
)

doc.add_heading("Status Reclassifications", level=3)

doc.add_paragraph(
    "4,537 loans were classified as \"Current\" but had zero outstanding principal balance, "
    "meaning they had been fully repaid. These were reclassified to \"Fully Paid\" to prevent "
    "them from inflating the active portfolio and distorting cash flow projections. "
    "Similarly, 168 delinquent loans (107 In Grace Period, 35 Late 16\u201330, 26 Late 31\u2013120) "
    "also had zero balances and were reclassified to Fully Paid, with boolean flags preserving "
    "their original status for transition analysis."
)

doc.add_heading("Late Fee and DTI Cleaning", level=3)

doc.add_paragraph(
    "Lending Club\u2019s minimum late fee is $15. We found 675 loans with late fees between "
    "$0 and $15 (likely rounding artifacts or partial reversals) and 8 loans with negative "
    "late fees. All were zeroed out. For DTI, one negative value was set to zero and four "
    "null values (joint applications missing both individual and joint DTI) were dropped."
)

doc.add_heading("Interest Rate Format", level=3)

doc.add_paragraph(
    "Interest rates were stored as whole-number percentages (e.g., 10.78) rather than "
    "decimals (0.1078). All downstream formulas expect decimal form, so a division by 100 "
    "was applied globally during the cleaning pipeline."
)

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 2 — METRIC AND CASH FLOW DEFINITIONS
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading("2. Metric and Cash Flow Definitions", level=2)

doc.add_heading("Key Pool-Level Metrics", level=3)

doc.add_paragraph(
    "All pool-level metrics are computed from the filtered cohort and are recalculated "
    "dynamically when the user changes the strata filter on the dashboard."
)

add_table(doc,
    ["Metric", "Definition", "Population"],
    [
        ["CDR (Annualized)", "1 \u2013 (1 \u2013 CDR_cum)^(12/WALA), where CDR_cum = "
         "defaulted UPB / total originated UPB", "All loans in cohort"],
        ["CPR", "1 \u2013 (1 \u2013 SMM)^12, where SMM = unscheduled principal / "
         "(beginning balance \u2013 scheduled principal)", "Current loans, March 2019 only"],
        ["Loss Severity", "Capped UPB lost / exposure, where exposure = "
         "funded_amnt \u2013 total_rec_prncp for Charged Off loans", "Charged Off loans"],
        ["WAC", "Balance-weighted average coupon (int_rate weighted by out_prncp)", "Current loans"],
        ["WAM", "Balance-weighted average remaining term (months)", "Current loans"],
    ],
)

doc.add_paragraph("")  # spacer

doc.add_paragraph(
    "An important methodological note: the raw CDR is a cumulative lifetime figure "
    "(8.73% for the full portfolio). Since the portfolio spans vintages from 2007 to 2018 with "
    "a weighted average loan age (WALA) of approximately 32 months, this cumulative rate must "
    "be annualized using a compound survival formula to produce a rate suitable for monthly "
    "cash flow projection. The annualized CDR for the full portfolio is 3.27%."
)

doc.add_heading("Monthly Cash Flow Projection", level=3)

doc.add_paragraph(
    "The projection engine operates at the pool-level aggregate: one balance, one WAC, "
    "one WAM, and one monthly payment per run. For each month t (starting April 2019):"
)

p = doc.add_paragraph(style="List Bullet")
p.add_run("Defaults").bold = True
p.add_run(" are applied first: defaults = beginning_balance \u00d7 MDR, "
          "split into losses (defaults \u00d7 loss severity) and recoveries.")

p = doc.add_paragraph(style="List Bullet")
p.add_run("Interest").bold = True
p.add_run(" accrues on the performing balance (after removing defaults): "
          "interest = performing_balance \u00d7 WAC/12.")

p = doc.add_paragraph(style="List Bullet")
p.add_run("Scheduled principal").bold = True
p.add_run(" follows standard amortization: monthly_payment \u2013 interest, "
          "capped at the performing balance.")

p = doc.add_paragraph(style="List Bullet")
p.add_run("Prepayments").bold = True
p.add_run(" are applied to the remaining balance: "
          "(performing_balance \u2013 scheduled_principal) \u00d7 SMM.")

p = doc.add_paragraph(style="List Bullet")
p.add_run("Total cash flow").bold = True
p.add_run(" to the investor = interest + total principal + recovery.")

doc.add_paragraph("")

doc.add_paragraph(
    "IRR is computed as the monthly discount rate that sets the NPV of all cash flows "
    "(with month-0 outlay of \u2013purchase_price \u00d7 UPB) to zero, then annualized via "
    "compounding: (1 + monthly_irr)^12 \u2013 1. A price solver (Brent\u2019s method) inverts "
    "this relationship to find the purchase price that achieves any target IRR."
)

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 3 — KEY INSIGHTS
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading("3. Key Insights About the Portfolio", level=2)

doc.add_heading("Portfolio Composition", level=3)

doc.add_paragraph(
    "The 2.26 million-loan portfolio has $33.9 billion in total originated UPB, of which "
    "$9.5 billion remains outstanding across 907,755 active loans (821,602 Current plus ~86,000 in delinquency stages). The portfolio is heavily "
    "concentrated in investment-grade credits: Grades A through C account for 77% of loans "
    "and 74% of originated volume. The weighted average coupon is 13.12%, with a weighted "
    "average remaining maturity of 34 months."
)

doc.add_heading("Credit Performance and Default Gradient", level=3)

doc.add_paragraph(
    "The annualized CDR exhibits a steep credit gradient across Lending Club grades:"
)

add_table(doc,
    ["Grade", "Annualized CDR", "Avg Rate", "Avg FICO"],
    [
        ["A", "0.73%", "7.08%", "~745"],
        ["B", "1.78%", "10.68%", "~710"],
        ["C", "3.40%", "14.14%", "~690"],
        ["D", "5.14%", "18.15%", "~680"],
        ["E", "7.07%", "21.85%", "~675"],
        ["F", "9.30%", "25.50%", "~670"],
        ["G", "11.35%", "28.21%", "~665"],
    ],
)

doc.add_paragraph("")

doc.add_paragraph(
    "The default rate roughly doubles with each two-grade step down: Grade A defaults at "
    "under 1% annually while Grade G defaults at over 11%. However, coupon rates also widen "
    "significantly (7% to 28%), meaning the net return question is whether the higher yield "
    "adequately compensates for the elevated credit risk. The pool-level loss severity is "
    "89.2% with a recovery rate of only 10.8%, reflecting the unsecured nature of these "
    "consumer loans\u2014there is no collateral to liquidate, so recovery depends entirely on "
    "post-charge-off collections."
)

doc.add_heading("Delinquency Pipeline", level=3)

doc.add_paragraph(
    "The transition matrix reveals that 1.52% of the Current loan pool enters the delinquency "
    "pipeline each period (moving to In Grace Period). Of those entering grace, roughly 60% "
    "remain in grace and 40% progress to Late (16\u201330 days). From the Late (16\u201330) bucket, "
    "about 23% cure back to Current, 44% remain in the same bucket, and 33% progress further "
    "to Late (31\u2013120 days). From Late (31\u2013120), approximately 68% remain and 32% charge off. "
    "This cascading pattern means that once a loan enters the delinquency pipeline, it has a "
    "substantial probability of eventually charging off\u2014the cure rates at each stage are "
    "relatively modest."
)

doc.add_heading("Prepayment Behavior", level=3)

doc.add_paragraph(
    "The observed CPR of 1.49% is low compared to typical structured finance benchmarks, "
    "reflecting the fact that Lending Club borrowers face limited refinancing options for "
    "unsecured personal loans. This low prepayment rate is favorable for investors purchasing "
    "at a discount, as it extends the period over which above-market coupons are collected."
)

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 4 — STRESS TESTING RECOMMENDATIONS
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading("4. Stress Testing Recommendations", level=2)

doc.add_paragraph(
    "The dashboard implements a multiplicative scenario framework where CDR and CPR are "
    "shifted by a user-adjustable percentage (default \u00b115%), while loss severity remains "
    "fixed. Below are the scenarios I would recommend and the rationale for each."
)

doc.add_heading("Recommended Scenario Set", level=3)

add_table(doc,
    ["Scenario", "CDR Shift", "CPR Shift", "Rationale"],
    [
        ["Base Case", "0%", "0%",
         "Historical performance of the filtered cohort. Assumes current "
         "economic conditions persist."],
        ["Mild Stress (+15%)", "+15%", "\u221215%",
         "A moderate deterioration: defaults rise 15%, prepayments slow 15%. "
         "Represents a garden-variety economic slowdown."],
        ["Severe Stress (+40%)", "+40%", "\u221230%",
         "A recessionary scenario. Unemployment spikes, borrowers lose ability "
         "to service debt, and refinancing activity dries up."],
        ["Upside (\u221215%)", "\u221215%", "+15%",
         "An improving economy: borrowers perform better and prepay more freely. "
         "IRR improves but WAL shortens."],
    ],
)

doc.add_paragraph("")

doc.add_heading("Why These Scenarios", level=3)

p = doc.add_paragraph()
p.add_run("Mild Stress (\u00b115%)").bold = True
p.add_run(
    " is the minimum sensitivity every investor should examine. It answers the question: "
    "\"If things get modestly worse, does this portfolio still meet my return threshold?\" "
    "For the full portfolio at a 95-cent purchase price, the base IRR is positive and "
    "remains positive under mild stress, confirming basic resilience."
)

p = doc.add_paragraph()
p.add_run("Severe Stress (+40% CDR / \u221230% CPR)").bold = True
p.add_run(
    " is calibrated to historical recession experience. During the 2008\u20132010 financial crisis, "
    "consumer unsecured credit defaults roughly doubled from pre-crisis levels. A +40% CDR "
    "shift approximates this magnitude for already-seasoned cohorts and reveals whether "
    "the coupon income can absorb significantly elevated losses. The CPR reduction to \u221230% "
    "reflects credit tightening: borrowers who would normally refinance cannot obtain new credit. "
    "This is the scenario an investor should focus on for downside protection\u2014if the portfolio "
    "still generates a positive IRR at this stress level, it has meaningful credit cushion."
)

p = doc.add_paragraph()
p.add_run("Upside (\u221215%)").bold = True
p.add_run(
    " is included to bound the range of outcomes. It shows the return potential if economic "
    "conditions improve, but also highlights a trade-off: higher prepayments reduce WAL and "
    "may cause the investor to be repaid faster than expected, creating reinvestment risk."
)

doc.add_heading("Additional Considerations", level=3)

doc.add_paragraph(
    "Loss severity is held fixed across all scenarios at 89.2%. This is conservative: in a true "
    "stress environment, recoveries may decline further as collection agencies face higher "
    "volumes and lower success rates. An investor could extend the framework by also stressing "
    "loss severity (e.g., 92\u201395% in severe stress) for a more pessimistic bound."
)

doc.add_paragraph(
    "Grade-level analysis is particularly valuable for stress testing. Grade A loans default at "
    "0.73% annually while Grade G defaults at 11.35%. An investor purchasing a slice of the "
    "portfolio could use the dashboard\u2019s strata filter to run scenarios on individual grades, "
    "identifying which credit tiers offer the best risk-adjusted returns under stress. For "
    "example, Grade C loans offer a 14.1% coupon against a 3.4% annualized CDR\u2014a ratio that "
    "may still be attractive even at +40% stress (CDR of 4.8%), whereas Grade F at 25.5% coupon "
    "versus 9.3% CDR has a thinner margin that erodes quickly under stress."
)

doc.add_paragraph(
    "Finally, vintage-level stress is important because earlier vintages (2012\u20132015) have fully "
    "seasoned and their default behavior is largely realized, while 2017\u20132018 vintages are "
    "still in their peak default window. Applying uniform stress across all vintages is "
    "conservative for the older cohorts and realistic for the newer ones."
)

# ═══════════════════════════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════════════════════════
output_path = "docs/portfolio_analysis_writeup.docx"
doc.save(output_path)
print(f"Saved to {output_path}")
